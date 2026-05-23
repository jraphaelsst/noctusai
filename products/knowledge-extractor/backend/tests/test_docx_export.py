"""docx_export — markdown→docx conversion + methodology/tree export."""
from docx import Document

from app.services.docx_export import (
    export_methodology_docx,
    export_tree_docx,
    md_to_docx,
)


def test_md_to_docx_creates_valid_docx(tmp_path):
    md = (
        "# Title\n\n## Section\n\n- a bullet with **bold**\n\n"
        "1. first item\n\nA plain paragraph.\n\n> a quote\n\n---\n"
    )
    out = tmp_path / "x.docx"
    md_to_docx(md, out)

    assert out.exists()
    doc = Document(str(out))
    texts = [p.text for p in doc.paragraphs]
    assert "Title" in texts
    assert any("bullet with bold" in t for t in texts)
    # the **bold** span became a bold run
    assert any(any(r.bold for r in p.runs) for p in doc.paragraphs)


def test_export_methodology_docx(tmp_path):
    (tmp_path / "METHODOLOGY.md").write_text("# Manual\n\nconteúdo", encoding="utf-8")
    (tmp_path / "methodology").mkdir()
    (tmp_path / "methodology" / "Mod 1.md").write_text("# Mod 1\n\nx", encoding="utf-8")

    written = export_methodology_docx(tmp_path)

    assert (tmp_path / "methodology" / "docx" / "METHODOLOGY.docx").exists()
    assert (tmp_path / "methodology" / "docx" / "Mod 1.docx").exists()
    assert len(written) == 2


def test_export_tree_docx_mirrors_per_module(tmp_path):
    root = tmp_path / "summaries"
    (root / "1 Mod A").mkdir(parents=True)
    (root / "1 Mod A" / "aula 1.md").write_text("# Aula 1\n\nx", encoding="utf-8")
    (root / "1 Mod A" / "aula 2.md").write_text("# Aula 2\n\ny", encoding="utf-8")
    (root / "2 Mod B").mkdir(parents=True)
    (root / "2 Mod B" / "aula 1.md").write_text("# B Aula 1\n\nz", encoding="utf-8")

    written = export_tree_docx(root)

    assert (root / "1 Mod A" / "docx" / "aula 1.docx").exists()
    assert (root / "1 Mod A" / "docx" / "aula 2.docx").exists()
    assert (root / "2 Mod B" / "docx" / "aula 1.docx").exists()
    assert len(written) == 3


def test_export_tree_docx_missing_root_is_noop(tmp_path):
    assert export_tree_docx(tmp_path / "nope") == []
