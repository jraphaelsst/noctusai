"""Export the methodology markdown files to .docx (via python-docx).

A focused markdown→docx converter for the structure these docs use: headings
(#/##/###), bullet and numbered lists, blockquotes, horizontal rules, and inline
**bold**. Not a general markdown engine — just what the methodology emits.
"""
from __future__ import annotations

import re
from pathlib import Path

_BOLD = re.compile(r"\*\*(.+?)\*\*")
_HEADING = re.compile(r"^(#{1,6})\s+(.*)$")
_BULLET = re.compile(r"^\s*[-*]\s+(.*)$")
_NUMBERED = re.compile(r"^\s*\d+\.\s+(.*)$")


def _add_runs(paragraph, text: str) -> None:
    """Add text to a paragraph, rendering **bold** spans as bold runs."""
    pos = 0
    for m in _BOLD.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        paragraph.add_run(m.group(1)).bold = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def md_to_docx(md_text: str, out_path: Path) -> Path:
    """Convert methodology markdown text to a .docx at out_path."""
    from docx import Document  # lazy

    doc = Document()
    for raw in md_text.splitlines():
        line = raw.rstrip()
        if not line.strip():
            continue
        if line.strip() == "---":
            continue  # horizontal rule → skip

        h = _HEADING.match(line)
        if h:
            level = min(len(h.group(1)), 4)
            doc.add_heading(_BOLD.sub(r"\1", h.group(2)).strip(), level=level)
            continue

        b = _BULLET.match(line)
        if b:
            _add_runs(doc.add_paragraph(style="List Bullet"), b.group(1))
            continue

        n = _NUMBERED.match(line)
        if n:
            _add_runs(doc.add_paragraph(style="List Number"), n.group(1))
            continue

        if line.lstrip().startswith(">"):
            _add_runs(doc.add_paragraph(style="Intense Quote"), line.lstrip()[1:].strip())
            continue

        _add_runs(doc.add_paragraph(), line)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path


def export_methodology_docx(data_dir: Path) -> list[Path]:
    """Convert METHODOLOGY.md + every data/methodology/*.md → methodology/docx/*.docx."""
    methodology_dir = data_dir / "methodology"
    out_dir = methodology_dir / "docx"

    sources: list[Path] = []
    master = data_dir / "METHODOLOGY.md"
    if master.exists():
        sources.append(master)
    if methodology_dir.exists():
        sources.extend(sorted(methodology_dir.glob("*.md")))

    written: list[Path] = []
    for src in sources:
        out = out_dir / f"{src.stem}.docx"
        md_to_docx(src.read_text(encoding="utf-8"), out)
        written.append(out)
    return written


def export_tree_docx(root: Path) -> list[Path]:
    """Convert every ``<module>/*.md`` under ``root`` into ``<module>/docx/*.docx``.

    Applies the methodology docx pattern (a ``docx/`` subfolder holding one .docx
    per source .md) to the per-module trees under ``data/summaries`` /
    ``data/transcripts``. Returns the .docx paths written. A no-op if ``root`` is
    missing or has no module subfolders with markdown.
    """
    written: list[Path] = []
    if not root.exists():
        return written
    for module_dir in sorted(p for p in root.iterdir() if p.is_dir()):
        if module_dir.name == "docx":
            continue  # never recurse into our own output folder
        out_dir = module_dir / "docx"
        for src in sorted(module_dir.glob("*.md")):
            out = out_dir / f"{src.stem}.docx"
            md_to_docx(src.read_text(encoding="utf-8"), out)
            written.append(out)
    return written


__all__ = ["md_to_docx", "export_methodology_docx", "export_tree_docx"]
